from __future__ import annotations

import shutil
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
WORK_DIR = REPO_ROOT / "build" / "manual_work"
SCREENSHOT_DIR = WORK_DIR / "screenshots"
CROP_DIR = WORK_DIR / "crops"
DOCS_DIR = REPO_ROOT / "docs"
DIST_MANUAL_DIR = REPO_ROOT / "dist" / "配布用_20260904" / "DigitalBuileder_GR" / "マニュアル"
OUTPUT_DOCX = DOCS_DIR / "電子請求書管理_操作マニュアル.docx"
DIST_DOCX = DIST_MANUAL_DIR / OUTPUT_DOCX.name
LOGO_PATH = REPO_ROOT / "assets" / "DigitalBuileder_GR.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION_FILL = "FFF8E8"
CAUTION_TEXT = "7A5A00"
WHITE = "FFFFFF"
TEAL = "009C9A"
ORANGE = "F05A00"
BODY_FONT = "Calibri"
EAST_ASIA_FONT = "Yu Gothic"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), BODY_FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), BODY_FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border(paragraph, color: str = BLUE, size: int = 12, space: int = 5) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title_style = doc.styles["Title"]
    title_style.font.color.rgb = RGBColor.from_string("000000")
    title_border = title_style._element.get_or_add_pPr().find(qn("w:pBdr"))
    if title_border is not None:
        title_style._element.get_or_add_pPr().remove(title_border)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "000000", 18, 10),
        ("Heading 2", 13, "000000", 14, 7),
        ("Heading 3", 12, "000000", 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = BODY_FONT
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_numbered_steps(doc: Document, steps) -> None:
    for index, text in enumerate(steps, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.first_line_indent = Inches(-0.22)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        set_run_font(paragraph.add_run(f"{index}. "), bold=True, color=NAVY)
        set_run_font(paragraph.add_run(text))


def configure_builtin_bullet(doc: Document) -> None:
    style = doc.styles["List Bullet"]
    style.font.name = BODY_FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25
    num_id_node = style._element.find("w:pPr/w:numPr/w:numId", style._element.nsmap)
    if num_id_node is None:
        return
    num_id = num_id_node.get(qn("w:val"))
    numbering = doc.part.numbering_part.element
    num = numbering.find(f"w:num[@w:numId='{num_id}']", numbering.nsmap)
    if num is None:
        return
    abstract_id = num.find("w:abstractNumId", numbering.nsmap).get(qn("w:val"))
    level = numbering.find(
        f"w:abstractNum[@w:abstractNumId='{abstract_id}']/w:lvl[@w:ilvl='0']",
        numbering.nsmap,
    )
    if level is None:
        return
    p_pr = level.find("w:pPr", numbering.nsmap)
    tabs = p_pr.find("w:tabs", numbering.nsmap)
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    tab = tabs.find("w:tab", numbering.nsmap)
    if tab is None:
        tab = OxmlElement("w:tab")
        tabs.append(tab)
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    ind = p_pr.find("w:ind", numbering.nsmap)
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    spacing = p_pr.find("w:spacing", numbering.nsmap)
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")


def add_bullet_item(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run_font(paragraph.add_run(text))


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_run_font(lead, bold=True, color=NAVY)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        set_run_font(p.add_run(text))


def add_callout(doc: Document, title: str, body: str, caution: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CAUTION_FILL if caution else CALLOUT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{title}  ")
    set_run_font(r1, bold=True, color=CAUTION_TEXT if caution else DARK_BLUE)
    r2 = p.add_run(body)
    set_run_font(r2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, bold=True, color=NAVY)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=9.5)
    set_table_geometry(table, widths_dxa)


def add_checklist(doc: Document, items: tuple[str, ...]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for text in items:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], CALLOUT)
        marker = cells[0].paragraphs[0]
        marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        marker.paragraph_format.space_after = Pt(0)
        set_run_font(marker.add_run("✓"), size=11, color=TEAL, bold=True)
        detail = cells[1].paragraphs[0]
        detail.paragraph_format.space_after = Pt(0)
        set_run_font(detail.add_run(text), size=10)
    set_table_geometry(table, [540, 8820])


def add_numbered_checklist(doc: Document, items: tuple[str, ...]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for index, text in enumerate(items, start=1):
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], LIGHT_BLUE)
        marker = cells[0].paragraphs[0]
        marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        marker.paragraph_format.space_after = Pt(0)
        set_run_font(marker.add_run(str(index)), size=10, color=NAVY, bold=True)
        detail = cells[1].paragraphs[0]
        detail.paragraph_format.space_after = Pt(0)
        set_run_font(detail.add_run(text), size=10)
    set_table_geometry(table, [540, 8820])


def add_picture(doc: Document, path: Path, width_inches: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cp = doc.add_paragraph(style="Caption")
    set_run_font(cp.add_run(caption), size=8.5, color=MUTED)


def crop_images() -> dict[str, Path]:
    CROP_DIR.mkdir(parents=True, exist_ok=True)
    jobs = {
        "list_main": ("02_invoice_list_tax_excluded.png", (0, 0, 1120, 245)),
        "tax_excluded": ("02_invoice_list_tax_excluded.png", (635, 35, 975, 225)),
        "tax_included": ("02b_invoice_list_tax_included.png", (635, 35, 975, 225)),
        "detail_top": ("03_invoice_detail.png", (0, 0, 1035, 240)),
        "pdf_marks": ("03_invoice_detail.png", (330, 240, 1135, 800)),
    }
    outputs: dict[str, Path] = {}
    for name, (source_name, box) in jobs.items():
        source = Image.open(SCREENSHOT_DIR / source_name)
        output = CROP_DIR / f"{name}.png"
        source.crop(box).save(output)
        outputs[name] = output
    return outputs


def set_headers_and_footers(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    set_run_font(header_p.add_run("電子請求書管理  操作マニュアル"), size=9, color=MUTED, bold=True)
    set_run_font(header_p.add_run("                                      DigitalBuileder_GR"), size=9, color=MUTED)
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    set_run_font(footer_p.add_run("Page "), size=9, color=MUTED)
    add_field(footer_p, "PAGE")
    set_run_font(footer_p.add_run(" / "), size=9, color=MUTED)
    add_field(footer_p, "NUMPAGES")


def add_page_title(doc: Document, number: str, title: str, subtitle: str = "") -> None:
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run(f"{number}  {title}")
    set_run_font(r, size=16, color="000000", bold=True)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(8)
        set_run_font(sp.add_run(subtitle), size=10, color=MUTED)


def add_page_break(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True


def build_manual() -> None:
    crops = crop_images()
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    DIST_MANUAL_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    set_headers_and_footers(doc)

    # Cover: editorial_cover pattern with a single branded title stack.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run("運用ユーザー向け"), size=11, color=TEAL, bold=True)
    add_picture(doc, LOGO_PATH, 2.35, "")
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("電子請求書管理\n操作マニュアル"), size=28, color="000000", bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    set_run_font(subtitle.add_run("DigitalBuileder_GR"), size=15, color=BLUE, bold=True)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    set_run_font(meta.add_run("2026年9月版"), size=11, color=NAVY, bold=True)
    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(scope.add_run("CSV＋zip取込 / 取込履歴 / 請求一覧 / 税表示切替 / 工種振分 / 工事表示 / PDF確認 / Excel出力"), size=10, color=MUTED)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(30)
    set_run_font(note.add_run("※掲載画面の会社名・金額・PDFは、すべて架空のサンプルです。"), size=9, color=MUTED, italic=True)

    add_page_break(doc)
    add_page_title(doc, "1", "まず知っておくこと", "金額表示と日常運用の基本ルール")
    add_callout(
        doc,
        "最重要",
        "CSVの請求金額(税込)は原本額として保持し、アプリ内の計算と工種振分入力は税抜を基準にします。初期表示は税抜で、［金額表示］から税込へ切り替えられます。",
        caution=True,
    )
    doc.add_paragraph("このアプリでできること", style="Heading 2")
    add_checklist(
        doc,
        (
            "Digital Billderから出力したCSVと、請求書PDFをまとめたzipをローカルへ取り込みます。",
            "請求一覧を工事・請求月・取引先・請求日で絞り込み、詳細・PDF・取込履歴を確認します。",
            "請求金額を工種コード別に振り分け、PDFへ確認用マークを配置します。",
            "完了工事はアーカイブして、自動処理と変更の対象から外します。",
            "月別の請求一覧と工種振分をExcelへ出力します。",
        ),
    )
    doc.add_paragraph("基本の作業順", style="Heading 2")
    workflow_steps = (
        "アプリを起動し、請求一覧を開く",
        "管理メニューからCSV＋zipをプレビューして取り込む",
        "請求一覧で対象を絞り込み、詳細を開く",
        "工種コード別振分とPDFマークを登録する",
        "必要に応じて確認用PDFまたはExcelを出力する",
    )
    add_numbered_checklist(doc, workflow_steps)
    add_table(
        doc,
        ["項目", "現在の扱い"],
        [
            ["消費税率", "10%固定（試験運用）"],
            ["税抜換算", "税込額 × 100 ÷ 110 の1円未満切り捨て"],
            ["初期表示", "税抜。切替後の状態は保存されます"],
        ],
        [2700, 6660],
    )

    add_page_break(doc)
    add_page_title(doc, "2", "起動する", "通常運用は起動.batから開始します")
    add_picture(doc, LOGO_PATH, 1.25, "アプリのGRアイコン")
    add_numbered_checklist(
        doc,
        (
            "「DigitalBuileder_GR」フォルダを開きます。",
            "起動.bat をダブルクリックします。デスクトップに「電子請求書管理」ショートカットがある場合は、そちらから起動できます。",
            "最初に［請求一覧］画面が開いたことを確認します。",
        ),
    )
    add_callout(
        doc,
        "フォルダ注意",
        "起動.bat、app.py、invoice_manager、assets、dataの配置を変えずに使用します。現行運用はEXEではなく、起動.batから開始します。",
        caution=True,
    )
    doc.add_paragraph("起動できないとき", style="Heading 2")
    add_table(
        doc,
        ["症状", "確認すること"],
        [
            ["起動.batを開いても起動しない", "Python環境とフォルダ内のapp.pyを確認"],
            ["Windowsの警告が出る", "社内管理者へ確認し、許可された配布物か確認"],
            ["ショートカットが開けない", "起動.batから直接起動し、管理者へ連絡"],
        ],
        [3000, 6360],
    )

    add_page_break(doc)
    add_page_title(doc, "3", "CSV＋zipを取り込む", "先にプレビューし、件数とエラーを確認します")
    add_picture(doc, SCREENSHOT_DIR / "01_csv_zip_import.png", 5.15, "CSV＋zip取込画面（サンプル）")
    add_numbered_steps(
        doc,
        (
            "請求一覧右上の［管理メニュー］を押し、［CSV＋zip取込］を開きます。",
            "［CSVファイル選択］と［zipファイル選択］で対象ファイルを指定します。ドラッグ＆ドロップも利用できます。",
            "請求月の自動判定を確認し、必要ならメモを入力します。",
            "［プレビュー］を押し、一致件数・新規件数・重複候補・エラー件数を確認します。CSVの必須列は［請求金額(税込)］で、画面の金額集計は税抜です。",
            "アーカイブ工事スキップ件数がある場合は、対象工事を確認します。アーカイブ中の工事は自動で表示へ戻りません。",
            "問題がなければ［取込実行］を押し、確認ダイアログで実行します。",
        ),
    )
    add_page_break(doc)
    add_page_title(doc, "4", "請求一覧を確認する", "絞り込み、並び替え、詳細確認をこの画面から行います")
    add_picture(doc, crops["list_main"], 6.15, "請求一覧上部（税抜表示のサンプル）")
    doc.add_paragraph("一覧でできること", style="Heading 2")
    add_table(
        doc,
        ["場所", "操作"],
        [
            ["上部フィルタ", "工事、請求月、取引先、請求日、並び順を指定"],
            ["一覧行", "クリックで選択。ダブルクリックで請求詳細を開く"],
            ["画面下部", "詳細、メモ、請求月変更、添付PDF、削除を実行"],
            ["右下", "表示件数と、現在の表示モードでの請求金額合計を確認"],
        ],
        [2500, 6860],
    )
    add_checklist(
        doc,
        (
            "行を選択すると、［詳細を開く］などのボタンが使用可能になります。",
            "複数行を選択して詳細を開くと、詳細画面の［前の請求］［次の請求］で移動できます。",
            "［選択工事の請求月再計算(試験)］は、工事を1件選んだ場合だけ実行できます。手動補正した請求月は変更しません。",
        ),
    )

    add_page_break(doc)
    add_page_title(doc, "5", "税抜／税込表示を切り替える", "切替は請求一覧と開いている請求詳細へ連動します")
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4680, 4680])
    for index, (label, image_path) in enumerate((("税抜表示", crops["tax_excluded"]), ("税込表示", crops["tax_included"]))):
        cell = table.cell(0, index)
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(label), size=10, color=NAVY, bold=True)
        r = p.add_run()
        r.add_picture(str(image_path), width=Inches(2.75))
    add_numbered_steps(
        doc,
        (
            "請求一覧上部の［金額表示］を開き、［税抜］または［税込］を選びます。",
            "一覧の見出し、各請求金額、右下の合計が選択した表示へ変わります。",
            "開いている請求詳細の請求金額と振分行も同じ表示へ変わります。振分合計、未振分額、超過額は常に税抜です。",
            "選択した表示は保存され、次回起動時にも引き継がれます。",
        ),
    )
    add_table(
        doc,
        ["税込保存額", "税抜表示", "計算"],
        [
            ["1,210,000円", "1,100,000円", "1,210,000 × 100 ÷ 110"],
            ["550,000円", "500,000円", "550,000 × 100 ÷ 110"],
            ["99,999円", "90,908円", "1円未満を切り捨て"],
        ],
        [2300, 2300, 4760],
    )
    add_callout(
        doc,
        "入力時の注意",
        "工種コード別振分の追加・編集は税抜で入力します。DBは税抜計算額と互換用の税込額を保持し、Excelは従来どおり税込で出力します。",
        caution=True,
    )

    add_page_break(doc)
    add_page_title(doc, "6", "請求詳細と工種コード別振分", "請求情報、添付PDF、振分状況を1画面で確認します")
    add_picture(doc, crops["detail_top"], 6.15, "請求詳細上部（税抜表示のサンプル）")
    doc.add_paragraph("工種コード別振分の登録", style="Heading 2")
    add_numbered_steps(
        doc,
        (
            "一覧から請求を選択し、［詳細を開く］または行をダブルクリックします。",
            "［振分行を追加］を押し、工種コードを選びます。",
            "金額を税抜で入力し、必要に応じてメモと並び順を設定して保存します。金額未確定なら空欄でも登録できます。",
            "修正は対象行を選んで［振分行を編集］、削除は［振分行を削除］を使います。",
        ),
    )
    add_callout(
        doc,
        "合計確認",
        "［請求金額(税抜)］［振分合計(税抜)］［未振分額］を確認します。振分が請求額を超えると［超過額］と警告が表示されます。",
    )
    add_body(doc, "工種コード行を選択すると、PDFプレビュー上部の［マーク対象］へ選択中のコードが表示されます。")

    add_page_break(doc)
    add_page_title(doc, "7", "PDFを確認し、マークを付ける", "原本PDFを変更せず、分類マーク付きPDFを別出力できます")
    add_picture(doc, crops["pdf_marks"], 3.50, "PDFプレビューと工種コードマーク（サンプル）")
    add_numbered_steps(
        doc,
        (
            "添付ファイル一覧でPDFを選択します。プレビューが自動表示されます。",
            "工種コード別振分の行を選択し、［マーク対象］を確認して［マーク配置］にチェックを付けます。",
            "PDF上をクリックしてマークを配置します。同じコードは複数箇所へ配置でき、［マーク一覧］で確認・削除できます。",
            "［確認用PDF出力］を押すと、マーク付きPDFが別ファイルで出力されます。",
        ),
    )
    add_table(
        doc,
        ["操作", "方法"],
        [
            ["ページ移動", "［前ページ］［次ページ］"],
            ["拡大・縮小", "［拡大］［縮小］、またはCtrl＋マウスホイール"],
            ["表示位置を移動", "PDF上を右ボタンでドラッグ"],
            ["直前のマークを戻す", "Ctrl＋Z"],
        ],
        [2600, 6760],
    )
    add_callout(doc, "原本保護", "画面上のマークはDBへ保存され、原本PDFには書き込みません。", caution=True)

    add_page_break(doc)
    add_page_title(doc, "8", "管理メニューと出力", "取込履歴、工事アーカイブ、マスタ管理、月別出力を行います")
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3600, 5760])
    left = table.cell(0, 0)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(SCREENSHOT_DIR / "04_management_menu.png"), width=Inches(2.45))
    right = table.cell(0, 1)
    set_cell_shading(right, LIGHT_GRAY)
    rp = right.paragraphs[0]
    rp.paragraph_format.space_after = Pt(4)
    set_run_font(rp.add_run("管理メニュー"), size=12, color=NAVY, bold=True)
    for label, detail in (
        ("CSV＋zip取込", "月次データの取込"),
        ("取込履歴", "取込件数・エラー・完了状態の確認"),
        ("工種コードマスタ", "工事ごとの工種コード管理"),
        ("取引先別工種候補", "振分候補の優先表示設定"),
        ("工事表示設定", "完了工事のアーカイブと再表示"),
        ("Digital Billderを開く", "登録済みURLをブラウザで開く"),
        ("Excel出力", "請求月を選んで月別一覧を出力"),
    ):
        pp = right.add_paragraph()
        pp.paragraph_format.space_after = Pt(3)
        set_run_font(pp.add_run(f"{label}: "), size=9.5, color=DARK_BLUE, bold=True)
        set_run_font(pp.add_run(detail), size=9.5)
    set_table_geometry(table, [3600, 5760])
    doc.add_paragraph("取込履歴と工事表示設定", style="Heading 2")
    add_body(doc, "［取込履歴］では、開始・完了日時、ファイル名、登録・PDF・エラー件数、状態を確認できます。［工事表示設定］では、工事コードまたは工事名で検索し、請求件数と最終請求日を確認できます。更新が終了した工事は、データを削除せずアーカイブできます。アーカイブ中は取込と変更処理の対象外です。")
    doc.add_paragraph("Excel出力", style="Heading 2")
    add_body(doc, "［Excel出力］で請求月を選んで出力し、完了メッセージに表示されたファイルを確認します。")
    add_callout(doc, "Excel金額", "Excelの請求金額と振分金額は税込です。画面の税抜／税込表示設定には連動しません。", caution=True)

    add_page_break(doc)
    add_page_title(doc, "9", "困ったとき・安全運用", "よくある確認ポイント")
    add_table(
        doc,
        ["症状", "対処"],
        [
            ["プレビューボタンが押せない", "CSVとzipの両方を指定します"],
            ["取込エラーがある", "エラー件数とメッセージを確認し、元データを修正して再度プレビューします"],
            ["金額が元CSVと違って見える", "［金額表示］が税抜か税込か確認します。元CSVは税込です"],
            ["新規・アーカイブ工事が選択肢にない", "取込画面と管理メニューを閉じます。アーカイブ工事は表示へ戻してから再取込します"],
            ["PDFが表示されない", "添付ファイルが登録されているか、対象行が正しいか確認します"],
            ["工種コードを選べない", "管理メニューの工種コードマスタを管理担当者が確認します"],
            ["Excelの金額が税抜表示と違う", "Excelは税込保存額を出力する仕様です"],
        ],
        [3100, 6260],
    )
    doc.add_paragraph("安全運用のチェック", style="Heading 2")
    add_checklist(
        doc,
        (
            "起動.bat、app.py、invoice_manager、assets、dataの配置を変えません。",
            "請求情報、取引先情報、PDF原本をメールやGitなどへ不用意に共有しません。",
            "取込、請求月変更・再計算、請求削除の直前にDBが自動バックアップされます。削除対象も必ず確認します。",
            "不明な警告やエラーが出た場合は、画面を閉じる前にスクリーンショットを保存して管理者へ連絡します。",
            "終了するときは請求一覧右上の［×］を押します。保存ダイアログが開いている場合は、先に完了またはキャンセルします。",
        ),
    )
    doc.core_properties.title = "電子請求書管理 操作マニュアル"
    doc.core_properties.subject = "DigitalBuileder_GR 運用ユーザー向け操作マニュアル"
    doc.core_properties.author = "角文株式会社"
    doc.core_properties.keywords = "電子請求書, DigitalBuileder_GR, 操作マニュアル"
    doc.save(OUTPUT_DOCX)
    shutil.copy2(OUTPUT_DOCX, DIST_DOCX)
    print(OUTPUT_DOCX)
    print(DIST_DOCX)


if __name__ == "__main__":
    build_manual()
